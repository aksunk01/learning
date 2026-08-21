from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.document import Document as _Document

from .base import DocumentElement, DocumentParser, ParsedDocument


def iter_block_items(parent):
    if isinstance(parent, _Document):
        parent_element = parent.element.body
    else:
        raise TypeError("Unsupported paretn type")

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)

class DOCXParser(DocumentParser):

    def parse(self, file_path: str) -> ParsedDocument:
        document = Document(file_path)

        elements: list[DocumentElement] = []
        current_section: str | None = None

        table_index = 0

        for block in iter_block_items(document):
            if isinstance(block, Paragraph):
                text = block.text.strip()

                if not text:
                    continue

                style_name = block.style.name if block.style else ""

                if style_name == "Title" or style_name.startswith("Heading"):
                    current_section = text

                element = DocumentElement(
                    text=text,
                    section=current_section,
                    metadata=self._get_element_metadata(
                        file_path,
                        element_type='paragraph',
                        style=style_name
                    ),
                )

                elements.append(element)

            elif isinstance(block, Table):
                for row_index, row in enumerate(block.rows):
                    cell_texts: list[str] = []

                    for cell in row.cells:
                        cell_text = " ".join(
                            paragraph.text.strip()
                            for paragraph in cell.paragraphs
                            if paragraph.text.strip()
                        )

                        if cell_text:
                            cell_texts.append(cell_text)

                    if not cell_texts:
                        continue

                    row_text = " | ".join(cell_texts)

                    element = DocumentElement(
                        text = row_text,
                        section=current_section,
                        metadata=self._get_element_metadata(
                            file_path,
                            element_type="table_row",
                            table_index=table_index,
                            row_index=row_index
                        )
                    )

                    elements.append(element)

                table_index +=1
        return ParsedDocument(
            elements=elements,
            metadata=self._get_file_metadata(
                file_path,
                "docx"
            )
        )